"""
Word 专业文档渲染器 —— JSON 结构 → 专业 docx。

设计哲学（LLM 出内容，程序出格式）:
    LLM 提供结构化 JSON（标题/段落/列表/表格），本模块负责全部格式：
    标题层级、中文字体(eastAsia)、首行缩进、目录字段、页边距等。

文档 JSON Schema:
    {
      "title": "文档标题",
      "subtitle": "可选副标题",
      "meta": [["文档编号", "ZP-RL-006"], ["版本号", "V3.0"]],   // 文档信息区（表格）
      "toc": true,                                                 // 是否插入目录
      "sections": [                                                // 章节（可嵌套）
        {
          "level": 1, "title": "1.0 目的总则",
          "paragraphs": ["正文段落..."],
          "bullets": ["列表项..."],                                 // 无序列表
          "numbered": ["编号项..."],                               // 有序列表
          "table": {"headers": ["列1", "列2"], "rows": [["a","b"]]},
          "sections": [ ... ]                                       // 子章节
        }
      ]
    }

用法:
    from adapters.word_renderer import render_document_json
    render_document_json(json_dict, "output.docx")
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 中文正文样式
BODY_FONT = "宋体"
BODY_FONT_EN = "Times New Roman"
BODY_SIZE = 12          # 小四
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)  # 深蓝标题
META_LABEL_COLOR = RGBColor(0x66, 0x66, 0x66)


def _set_run_font(run, name_cn=BODY_FONT, name_en=BODY_FONT_EN, size=BODY_SIZE,
                  bold=False, color=None):
    """同时设置中西文字体（关键：eastAsia 缺省会字体混乱）。"""
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)


def _add_toc_field(doc):
    """插入 Word 目录字段（打开后按 F9 或右键更新）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "（目录：请在 Word 中右键此处 → 更新域）"
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)


def _add_paragraph(doc, text, indent_chars=2, align=None, space_after=6):
    """正文段落：首行缩进 2 字符 + 中文宋体 + 西文 Times New Roman。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent_chars:
        pf.first_line_indent = Pt(BODY_SIZE * indent_chars)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if align:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run)
    return p


def _add_heading(doc, text, level):
    """标题：用 Heading 样式（真结构，可导航）+ 统一字体。"""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        _set_run_font(run, name_cn="黑体", name_en="Arial",
                      size=18 if level == 1 else (15 if level == 2 else 13),
                      bold=True, color=HEADING_COLOR)
    return h


def _add_bullets(doc, items, numbered=False):
    """列表：左缩进 + 项目符号/编号。"""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Pt(BODY_SIZE * 2)
        pf.first_line_indent = Pt(-BODY_SIZE)  # 悬挂缩进
        pf.space_after = Pt(3)
        pf.line_spacing = 1.5
        prefix = f"{i}. " if numbered else "• "
        run = p.add_run(prefix + item)
        _set_run_font(run)


def _add_meta_table(doc, meta):
    """文档信息区：两列表格（标签左对齐灰字，值右对齐）。"""
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta):
        c0 = table.cell(i, 0)
        c0.width = Cm(4)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(str(label))
        _set_run_font(r0, size=10.5, color=META_LABEL_COLOR)
        c1 = table.cell(i, 1)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(str(value))
        _set_run_font(r1, size=10.5)
    doc.add_paragraph()  # 表后空行


def _add_table(doc, table_spec):
    """数据表格。"""
    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])
    t = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=len(headers) or 1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        for j, h in enumerate(headers):
            cell = t.cell(0, j)
            p = cell.paragraphs[0]
            r = p.add_run(str(h))
            _set_run_font(r, size=10.5, bold=True, color=HEADING_COLOR)
    for i, row in enumerate(rows, start=1 if headers else 0):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            _set_run_font(r, size=10.5)
    doc.add_paragraph()


def _render_section(doc, sec):
    """渲染一个章节（可嵌套子章节）。"""
    if sec.get("title"):
        _add_heading(doc, sec["title"], min(sec.get("level", 1), 4))
    for p in sec.get("paragraphs", []):
        _add_paragraph(doc, p)
    for b in sec.get("bullets", []):
        _add_bullets(doc, b if isinstance(b, list) else [b])
    for n in sec.get("numbered", []):
        _add_bullets(doc, n if isinstance(n, list) else [n], numbered=True)
    if sec.get("table"):
        _add_table(doc, sec["table"])
    for sub in sec.get("sections", []):
        _render_section(doc, sub)


def render_document_json(data: dict, output_path: str) -> str:
    """渲染文档 JSON 到 docx，返回成功消息。"""
    doc = Document()

    # 页面边距（公文标准：上下 2.54cm 左右 3.18cm 近似）
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

    # 标题
    if data.get("title"):
        t = doc.add_heading(data["title"], level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            _set_run_font(run, name_cn="黑体", name_en="Arial", size=22, bold=True,
                          color=HEADING_COLOR)
    if data.get("subtitle"):
        st = doc.add_paragraph()
        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = st.add_run(data["subtitle"])
        _set_run_font(r, size=12, color=META_LABEL_COLOR)

    # 文档信息区
    if data.get("meta"):
        _add_meta_table(doc, data["meta"])

    # 目录
    if data.get("toc"):
        toc_title = doc.add_heading("目录", level=1)
        for run in toc_title.runs:
            _set_run_font(run, name_cn="黑体", name_en="Arial", size=15, bold=True,
                          color=HEADING_COLOR)
        _add_toc_field(doc)
        doc.add_page_break()

    # 章节
    for sec in data.get("sections", []):
        _render_section(doc, sec)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return f"✅ 专业文档已生成: {out.name}（{len(data.get('sections', []))} 个章节）"


def render_document_json_str(json_str: str, output_path: str) -> str:
    """从 JSON 字符串渲染（供 LLM 工具调用）。"""
    try:
        data = json.loads(json_str)
    except json.JSONDecodeError as e:
        return f"❌ JSON 解析失败: {e}（请检查大括号/引号/逗号）"
    return render_document_json(data, output_path)
