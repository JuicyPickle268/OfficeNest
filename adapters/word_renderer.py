"""
Word 只读渲染器：python-docx 解析 → Qt 面板格式化显示。
文件被 Word 打开也不影响——只走 read_only 通道。
"""
from pathlib import Path
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QTextEdit, QLabel, QDialog, QPushButton, QHBoxLayout
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QFont, QTextCursor, QTextCharFormat, QColor, QTextTableFormat


def render_docx_to_qtextedit(filepath: str, text_edit: QTextEdit):
    """将 .docx 渲染到 QTextEdit 中（只读预览，处理文件锁）。"""
    text_edit.clear()
    path = Path(filepath)
    if not path.exists():
        text_edit.setText(f"文件不存在: {filepath}")
        return

    # 如果文件被锁定，先复制到临时目录再读
    read_path = str(path)
    tmp_path = None
    try:
        f = open(str(path), "a"); f.close()
    except (IOError, PermissionError):
        import tempfile, shutil
        tmp = Path(tempfile.gettempdir()) / f"mother_preview_{path.name}"
        shutil.copy2(str(path), str(tmp))
        read_path = str(tmp)
        tmp_path = str(tmp)

    try:
        from docx import Document
        doc = Document(read_path)
    except Exception as e:
        text_edit.setText(f"无法打开: {e}")
        return

    cursor = text_edit.textCursor()
    cursor.movePosition(QTextCursor.MoveOperation.End)

    for para in doc.paragraphs:
        fmt = QTextCharFormat()
        text = para.text

        # 样式映射
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name or "标题 1" in style_name:
            fmt.setFontPointSize(20)
            fmt.setFontWeight(QFont.Weight.Bold)
        elif "heading 2" in style_name or "标题 2" in style_name:
            fmt.setFontPointSize(16)
            fmt.setFontWeight(QFont.Weight.Bold)
        elif "heading 3" in style_name or "标题 3" in style_name:
            fmt.setFontPointSize(13)
            fmt.setFontWeight(QFont.Weight.Bold)
        else:
            fmt.setFontPointSize(10)

        if para.runs:
            r = para.runs[0]
            if r.font.bold:
                fmt.setFontWeight(QFont.Weight.Bold)
            if r.font.size:
                fmt.setFontPointSize(r.font.size.pt)

        cursor.insertText(text + "\n", fmt)

    # 表格
    for table in doc.tables:
        cursor.insertText("\n", QTextCharFormat())
        rows = len(table.rows)
        cols = len(table.columns)
        qt_table = cursor.insertTable(rows + 1, cols)

        # 表头格式
        header_fmt = QTextCharFormat()
        header_fmt.setFontWeight(QFont.Weight.Bold)
        header_fmt.setBackground(QColor("#4472C4"))
        header_fmt.setForeground(QColor("white"))

        cell_fmt = QTextCharFormat()
        cell_fmt.setFontPointSize(9)

        for ri in range(rows):
            for ci in range(cols):
                cell = table.cell(ri, ci)
                cell_cursor = qt_table.cellAt(ri, ci).firstCursorPosition()
                cell_cursor.insertText(
                    cell.text.strip()[:100],
                    header_fmt if ri == 0 else cell_fmt
                )

        cursor.movePosition(QTextCursor.MoveOperation.End)
        cursor.insertText("\n", QTextCharFormat())

    # 合并信息
    cursor.insertText("\n", QTextCharFormat())
    info_fmt = QTextCharFormat()
    info_fmt.setForeground(QColor("#808080"))
    info_fmt.setFontPointSize(8)
    cursor.insertText(
        f"[Word 预览] {filepath} | {len(doc.paragraphs)} 段落 | {len(doc.tables)} 表格\n",
        info_fmt
    )
    text_edit.setTextCursor(cursor)


class WordPreviewDialog(QDialog):
    """Word 预览弹窗。"""

    def __init__(self, filepath: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Word 预览: {Path(filepath).name}")
        self.resize(900, 650)

        layout = QVBoxLayout(self)
        self._text = QTextEdit()
        self._text.setReadOnly(True)
        self._text.setFont(QFont("Microsoft YaHei", 10))
        self._text.setStyleSheet("QTextEdit { background: white; color: #1a1a1a; }")
        layout.addWidget(self._text)

        btn = QPushButton("关闭")
        btn.clicked.connect(self.close)
        layout.addWidget(btn)

        render_docx_to_qtextedit(filepath, self._text)
