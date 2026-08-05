"""
Word Skill：5 个 LLM 可调用的 Word 工具。
"""
from pathlib import Path
from skills.base import BaseSkill
from adapters.office_bridge import OfficeBridge
from core.interfaces.office_bridge import TableData


class WordSkill(BaseSkill):

    def __init__(self, office: OfficeBridge):
        self._office = office

    @property
    def _has_win32(self) -> bool:
        try:
            import win32com.client
            return True
        except ImportError:
            return False

    @property
    def name(self) -> str:
        return "word"

    def get_tools(self) -> list[dict]:
        return [
            {"name": "word_generate", "fn": self.word_generate,
             "schema": self._schema("基于模板生成 Word 文档", {
                 "template_path": {"type": "string", "description": "模板文件路径"},
                 "output_path": {"type": "string", "description": "输出文件路径"},
                 "replacements": {"type": "object", "description": "占位符替换映射，如 {'{日期}': '2026-07-17', '{姓名}': '张三'}"},
             }, ["template_path", "output_path"])},
            {"name": "word_create_from_scratch", "fn": self.word_create_from_scratch,
             "schema": self._schema("从零创建 Word 文档", {
                 "output_path": {"type": "string", "description": "输出文件路径"},
                 "title": {"type": "string", "description": "文档标题"},
                 "content": {"type": "string", "description": "文档正文"},
             }, ["output_path", "content"])},
            {"name": "word_fill_table", "fn": self.word_fill_table,
             "schema": self._schema("填充 Word 中的表格", {
                 "filepath": {"type": "string", "description": "Word 文件路径"},
                 "table_index": {"type": "integer", "description": "表格序号，从1开始"},
                 "headers": {"type": "array", "items": {"type": "string"}, "description": "表头"},
                 "rows": {"type": "array", "description": "数据行"},
             }, ["filepath", "headers", "rows"])},
            {"name": "word_export_pdf", "fn": self.word_export_pdf,
             "schema": self._schema("将 Word 导出为 PDF", {
                 "filepath": {"type": "string", "description": "Word 文件路径"},
                 "pdf_path": {"type": "string", "description": "PDF 输出路径"},
             }, ["filepath", "pdf_path"])},
            {"name": "word_list_templates", "fn": self.word_list_templates,
             "schema": self._schema("列出所有可用的 Word 模板", {})},
            {"name": "word_read", "fn": self.word_read,
             "schema": self._schema("读取 Word 文件的纯文本内容", {
                 "filepath": {"type": "string", "description": "Word 文件路径"},
             }, ["filepath"])},
            {"name": "word_read_table", "fn": self.word_read_table,
             "schema": self._schema("读取 Word 文件中指定表格的结构（行列坐标+合并信息）", {
                 "filepath": {"type": "string", "description": "Word 文件路径"},
                 "table_index": {"type": "integer", "description": "表格序号，从1开始，默认1"},
             }, ["filepath"])},
            {"name": "word_batch_generate", "fn": self.word_batch_generate,
             "schema": self._schema("批量生成 Word：Excel 每行 → 填充模板 → 每人一份独立文档", {
                 "template_path": {"type": "string", "description": "Word 模板路径"},
                 "excel_path": {"type": "string", "description": "Excel 数据源路径"},
                 "sheet": {"type": "string", "description": "Excel Sheet名称"},
                 "output_dir": {"type": "string", "description": "输出目录，默认 ./output"},
                 "name_column": {"type": "string", "description": "用于文件命名的列"},
                 "field_map": {"type": "object", "description": "Excel列名→Word占位符映射"},
             }, ["template_path", "excel_path", "field_map"])},
            {"name": "excel_paste_to_word", "fn": self.excel_paste_to_word,
             "schema": self._schema("把 Excel 中的表格区域（含格式/合并/颜色）粘贴到 Word——Excel 建表，Word 排版", {
                 "excel_path": {"type": "string", "description": "Excel 文件路径"},
                 "excel_range": {"type": "string", "description": "要复制的区域如 A1:G20"},
                 "word_path": {"type": "string", "description": "Word 文件路径"},
                 "word_bookmark": {"type": "string", "description": "Word 书签名（可选），为空贴到文末"},
             }, ["excel_path", "excel_range", "word_path"])},
            {"name": "word_insert_image", "fn": self.word_insert_image,
             "schema": self._schema("向 Word 文档插入图片（图表截图、照片等）", {
                 "filepath": {"type": "string", "description": "Word 文件路径"},
                 "image_path": {"type": "string", "description": "图片文件路径（PNG/JPG）"},
                 "width_inches": {"type": "number", "description": "图片宽度（英寸），默认5.5"},
             }, ["filepath", "image_path"])},
            {"name": "excel_chart_to_word", "fn": self.excel_chart_to_word,
             "schema": self._schema("把 Excel 图表以可编辑对象嵌入 Word（和你手动 Ctrl+C/V 一样）", {
                 "excel_path": {"type": "string", "description": "Excel 文件路径（需在 Excel 中打开）"},
                 "sheet": {"type": "string", "description": "Sheet 名称"},
                 "chart_index": {"type": "integer", "description": "图表序号，1=第一个图表"},
                 "word_path": {"type": "string", "description": "Word 文件路径"},
             }, ["excel_path", "word_path"])},
            {"name": "word_append_paragraph", "fn": self.word_append_paragraph,
             "schema": self._schema("向 Word 文档末尾追加文字段落——实现表格/文字混排", {
                 "filepath": {"type": "string", "description": "Word 文件路径"},
                 "text": {"type": "string", "description": "段落文字内容"},
                 "style": {"type": "string", "description": "样式: normal/heading1/heading2，默认normal"},
             }, ["filepath", "text"])},
        ]

    async def word_generate(self, template_path: str, output_path: str,
                            replacements: dict | None = None) -> str:
        if not self._has_win32:
            return self._word_fallback_replace(template_path, output_path, replacements or {})
        sid = await self._office.word_open(template_path)
        try:
            if replacements:
                for placeholder, value in replacements.items():
                    await self._office.word_replace(sid, placeholder, str(value))

            await self._office.word_save_as(sid, output_path)
            return f"✅ 文档已生成: {Path(output_path).name}（替换了 {len(replacements or {})} 处占位符）"
        finally:
            await self._office.word_close(sid)

    async def word_create_from_scratch(self, output_path: str, content: str, title: str = "") -> str:
        if not self._has_win32:
            return self._word_fallback_create(output_path, title, content)
        try:
            return await self._word_create_win32(output_path, content, title)
        except Exception as e:
            # win32 失败自动降级到 python-docx
            return self._word_fallback_create(output_path, title, content)

    async def _word_create_win32(self, output_path: str, content: str, title: str = "") -> str:
        sid = await self._office.word_open(output_path)
        try:
            app = self._office._sessions[sid]["app"]
            doc = self._office._sessions[sid]["doc"]
            selection = app.Selection
            if title:
                selection.Style = 2  # wdStyleHeading1
                selection.TypeText(title)
                selection.TypeParagraph()
            # 正文用 Normal 样式，确保字体统一
            selection.Style = -1  # wdStyleNormal
            selection.Font.Name = "宋体"
            selection.Font.Size = 11
            selection.Font.Bold = False
            # 长文本分段写入
            for i in range(0, len(content), 2000):
                chunk = content[i:i + 2000]
                selection.TypeText(chunk)
            await self._office.word_save_as(sid, output_path)
            return f"✅ 文档已创建: {Path(output_path).name}"
        finally:
            await self._office.word_close(sid)

    async def word_fill_table(self, filepath: str, headers: list[str],
                               rows: list[list], table_index: int = 1) -> str:
        sid = await self._office.word_open(filepath)
        try:
            td = TableData(headers=headers, rows=rows)
            count = await self._office.word_fill_table(sid, table_index, td)
            await self._office.word_save_as(sid, filepath)
            return f"✅ 表格已填充: {count} 行数据"
        finally:
            await self._office.word_close(sid)

    async def word_export_pdf(self, filepath: str, pdf_path: str = "") -> str:
        if not pdf_path:
            pdf_path = str(Path(filepath).with_suffix(".pdf"))
        sid = await self._office.word_open(filepath)
        try:
            await self._office.word_refresh_fields(sid)
            ok = await self._office.word_export_pdf(sid, pdf_path)
            if ok:
                return f"✅ PDF 已导出: {Path(pdf_path).name}"
            else:
                return f"⚠️ PDF 导出失败（文件未生成），请手动用 Word 打开另存为 PDF。docx 文件仍在: {filepath}"
        finally:
            await self._office.word_close(sid)

    async def word_list_templates(self) -> str:
        from adapters.file_bridge import FileBridge
        fb = FileBridge()
        files = fb.list_dir("./templates", "*.docx") + fb.list_dir("./templates", "*.doc")
        if not files:
            return "（templates/ 目录下暂无模板文件）"
        return "\n".join(f"  {f.name}" for f in files)

    def word_read(self, filepath: str) -> str:
        try:
            path = self._ensure_docx(filepath)
        except (FileNotFoundError, RuntimeError) as e:
            return str(e)
        try:
            from docx import Document
            doc = Document(str(path))
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            # 也读表格
            for i, table in enumerate(doc.tables):
                lines.append(f"\n[表格{i+1}]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    lines.append(" | ".join(cells))
            return "\n".join(lines[:200]) if lines else "（文档为空）"
        except Exception as e:
            return f"❌ 读取失败: {e}"

    def word_preview(self, filepath: str) -> str:
        """打开 Word 只读预览窗口。"""
        try:
            path = self._ensure_docx(filepath)
        except Exception as e:
            return str(e)
        try:
            from PySide6.QtWidgets import QApplication
            from adapters.word_renderer import WordPreviewDialog
            app = QApplication.instance()
            if app:
                dlg = WordPreviewDialog(str(path))
                dlg.exec()
                return f"✅ Word 预览窗口已打开: {Path(filepath).name}"
            return "⚠️ 未检测到 Qt 应用实例"
        except Exception as e:
            return f"❌ 预览失败: {e}"

    def word_read_table(self, filepath: str, table_index: int = 1) -> str:
        try:
            path = self._ensure_docx(filepath)
        except (FileNotFoundError, RuntimeError) as e:
            return str(e)
        try:
            from docx import Document
            doc = Document(str(path))
            if table_index > len(doc.tables):
                return f"❌ 表格序号 {table_index} 超出范围（共 {len(doc.tables)} 个表格）"
            table = doc.tables[table_index - 1]
            lines = [f"表格 {table_index}: {len(table.rows)} 行 × {len(table.columns)} 列"]
            for ri, row in enumerate(table.rows):
                cells_info = []
                for ci, cell in enumerate(row.cells):
                    text = cell.text.strip().replace("\n", " ") if cell.text else ""
                    # 检查合并
                    tc = cell._tc
                    grid_span = int(tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan').get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1')) if tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan') is not None else 1
                    vmerge = tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge') is not None
                    note = ""
                    if grid_span > 1 or vmerge:
                        note = f" [合并{grid_span}列]" if grid_span > 1 else ""
                        note += " [垂直合并]" if vmerge else ""
                    display = f"({ri+1},{ci+1}){note}={text}" if text else f"({ri+1},{ci+1}){note}=(空)"
                    cells_info.append(display)
                lines.append("  " + "  |  ".join(cells_info))
            return "\n".join(lines[:100])
        except Exception as e:
            return f"❌ 解析失败: {e}"

    def word_batch_generate(self, template_path: str, excel_path: str, field_map: dict,
                             sheet: str = "", output_dir: str = "./output",
                             name_column: str = "姓名") -> str:
        """批量生成：Excel每行→填充模板→每人一份独立docx。"""
        import openpyxl
        path_t = Path(template_path)
        path_e = Path(excel_path)
        if not path_t.exists():
            path_t = Path("./templates") / path_t.name
        if not path_e.exists():
            path_e = Path("./workbooks") / path_e.name
        if not path_t.exists():
            return f"❌ 模板不存在: {template_path}"
        if not path_e.exists():
            return f"❌ Excel 不存在: {excel_path}"

        wb = openpyxl.load_workbook(str(path_e), read_only=True, data_only=True)
        ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active

        # 读表头行建列索引
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            wb.close(); return "❌ Excel 为空"
        headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(rows[0])]
        name_idx = headers.index(name_column) if name_column in headers else 0

        Path(output_dir).mkdir(parents=True, exist_ok=True)
        count = 0
        errors = []

        for row_idx, row_data in enumerate(rows[1:]):
            if not row_data or all(v is None or str(v).strip() == "" for v in row_data):
                continue
            # 构建替换字典
            replacements = {}
            for excel_col, placeholder in field_map.items():
                if excel_col in headers:
                    col_idx = headers.index(excel_col)
                    val = str(row_data[col_idx]) if col_idx < len(row_data) and row_data[col_idx] is not None else ""
                    replacements[placeholder] = val

            # 生成文件名
            name_val = str(row_data[name_idx]) if name_idx < len(row_data) and row_data[name_idx] else f"row_{row_idx+1}"
            safe_name = "".join(c if c.isalnum() or c in "._- " else "_" for c in name_val).strip()
            out_path = str(Path(output_dir) / f"{safe_name}_员工信息表.docx")

            try:
                from docx import Document
                doc = Document(str(path_t))
                # 替换段落
                for para in doc.paragraphs:
                    for key, val in replacements.items():
                        if key in para.text:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
                # 替换表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key, val in replacements.items():
                                if key in cell.text:
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            if key in run.text:
                                                run.text = run.text.replace(key, val)
                doc.save(out_path)
                count += 1
            except Exception as e:
                errors.append(f"{name_val}: {e}")

        wb.close()
        result = f"✅ 批量生成完成: {count} 份文档 → {output_dir}/"
        if errors:
            result += f"\n⚠️ {len(errors)} 个失败: {'; '.join(errors[:3])}"
        return result

    def word_insert_image(self, filepath: str, image_path: str, width_inches: float = 5.5) -> str:
        """向 Word 文档插入图片（python-docx）。"""
        from docx import Document
        from docx.shared import Inches
        path = Path(filepath)
        img = Path(image_path)
        if not path.is_absolute():
            path = Path("./output") / path.name
        if not img.exists():
            return f"❌ 图片不存在: {image_path}"
        try:
            doc = Document(str(path)) if path.exists() else Document()
            doc.add_picture(str(img), width=Inches(width_inches))
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
            return f"✅ 图片已插入: {Path(filepath).name}"
        except Exception as e:
            return f"❌ 插入失败: {e}"

    def word_append_paragraph(self, filepath: str, text: str, style: str = "normal") -> str:
        """向 Word 末尾追加文字段落——实现表格/文字混排。"""
        from docx import Document
        path = Path(filepath)
        if not path.is_absolute():
            for d in ["./output", "./workbooks", "."]:
                p = Path(d) / path.name
                if p.exists(): path = p; break
        try:
            doc = Document(str(path)) if path.exists() else Document()
            self._set_default_font(doc)
            if style == "heading1": doc.add_heading(text, level=1)
            elif style == "heading2": doc.add_heading(text, level=2)
            else: doc.add_paragraph(text)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
            return f"✅ 段落已追加: {path.name}"
        except Exception as e:
            return f"❌ 追加失败: {e}"

    def excel_chart_to_word(self, excel_path: str, word_path: str,
                             sheet: str = "", chart_index: int = 1) -> str:
        """Excel 图表以 OLE 对象嵌入 Word（不是截图，和你手动 Ctrl+C/V 一样）。"""
        import win32com.client

        # ── 找到 Excel 中的图表 ──
        try:
            excel = win32com.client.GetActiveObject("Excel.Application")
        except Exception:
            return "❌ Excel 未运行，请先打开 Excel 文件"

        path_e = Path(excel_path).resolve()
        wb = None
        for w in excel.Workbooks:
            try:
                if Path(w.FullName).resolve() == path_e:
                    wb = w; break
            except Exception:
                continue
        if not wb:
            return "❌ 文件未在 Excel 中打开"

        ws = wb.ActiveSheet
        if sheet:
            for s in wb.Worksheets:
                if s.Name == sheet:
                    ws = s; break

        # 尝试复制图表
        copied = False
        try:
            ws.ChartObjects(chart_index).Copy()
            copied = True
        except Exception:
            try:
                wb.Charts(chart_index).Copy()
                copied = True
            except Exception:
                pass

        if not copied:
            return f"❌ 找不到图表 #{chart_index}"

        # ── 粘贴到 Word ──
        try:
            word = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
            except Exception:
                return "❌ Word 不可用"

        path_w = Path(word_path).resolve()
        doc = None
        for d in word.Documents:
            try:
                if Path(d.FullName).resolve() == path_w:
                    doc = d; break
            except Exception:
                continue
        if not doc:
            if path_w.exists():
                doc = word.Documents.Open(str(path_w))
            else:
                doc = word.Documents.Add()

        doc.Activate()
        selection = word.Selection
        selection.EndKey(6)  # wdStory → 跳到文末
        selection.Paste()

        doc.Save()
        return f"✅ 图表已嵌入 Word（OLE对象，双击可编辑）: {Path(word_path).name}"

    def _set_default_font(self, doc, font_name="宋体", font_size=11):
        """设置 python-docx 文档的默认字体。"""
        from docx.shared import Pt
        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

    async def excel_paste_to_word(self, excel_path: str, excel_range: str,
                                    word_path: str, sheet: str = "", word_bookmark: str = "") -> str:
        """Excel 区域 → 粘贴到 Word（保留格式/合并/颜色）。"""
        # sheet 参数用于指定 Excel Sheet（透传给 OfficeBridge）
        import asyncio
        # 打开两端
        xl_sid = await self._office.excel_open(excel_path)
        wd_sid = await self._office.word_open(word_path)
        try:
            ok = await self._office.excel_range_copy_to_word(
                xl_sid, excel_range, wd_sid, word_bookmark
            )
            if ok:
                await self._office.word_save_as(wd_sid, word_path)
                return f"✅ Excel {excel_range} 已粘贴到 Word: {Path(word_path).name}"
            else:
                return "❌ 粘贴失败：请确保 Excel 和 Word 文件都已打开"
        finally:
            await self._office.excel_close(xl_sid)
            await self._office.word_close(wd_sid)

    # ── python-docx 降级方案 ──

    def _word_fallback_create(self, output_path: str, title: str, content: str) -> str:
        """pywin32 不可用时用 python-docx 创建纯文本文档。"""
        try:
            from docx import Document
            doc = Document()
            self._set_default_font(doc)
            if title:
                doc.add_heading(title, level=1)
            doc.add_paragraph(content)
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return f"✅ 文档已创建（python-docx降级模式）: {Path(output_path).name}"
        except Exception as e:
            return f"❌ Word创建失败: {e}"

    def _word_fallback_replace(self, template_path: str, output_path: str, replacements: dict) -> str:
        """pywin32 不可用时用 python-docx 替换占位符。"""
        try:
            from docx import Document
            doc = Document(template_path)
            count = 0
            for para in doc.paragraphs:
                for key, val in replacements.items():
                    if key in para.text:
                        para.text = para.text.replace(key, str(val))
                        count += 1
            # 也检查表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, val in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, str(val))
                                count += 1
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return f"✅ 文档已生成（python-docx降级模式）: {Path(output_path).name}（替换了 {count} 处占位符）"
        except Exception as e:
            return f"❌ Word生成失败: {e}"



    @staticmethod
    def _ensure_docx(filepath: str) -> Path:
        """确保文件是 .docx 可读的（.doc 自动转换）。"""
        original = filepath
        path = Path(filepath)
        if not path.is_absolute():
            for d in ["./workbooks", "./templates", "./output"]:
                p = Path(d) / path.name
                if p.exists():
                    path = p.resolve(); break
        if not path.exists():
            # 再试精确匹配
            for d in ["./workbooks", "./templates", "./output"]:
                for f in Path(d).glob("*"):
                    if path.name in f.name or f.name in path.name:
                        path = f.resolve(); break
                if path.exists():
                    break
        if not path.exists():
            raise FileNotFoundError(f"找不到文件: {original}（已搜索 workbooks/templates/output）")
        if path.suffix.lower() in ('.docx', '.docm'):
            return path
        if path.suffix.lower() == '.doc':
            # .doc 不支持——原样返回 .docx 路径提示
            docx_hint = str(path.with_suffix('.docx'))
            raise FileNotFoundError(
                f"不支持 .doc 格式。请用 Word 打开 → 另存为 .docx → 放回原目录。\n"
                f"期望路径: {docx_hint}"
            )
        return path
