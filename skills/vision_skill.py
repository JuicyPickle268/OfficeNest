"""
视觉 Skill：截图 Excel/Word → 多 provider 视觉分析（GLM优先，商汤备用）。
"""
import sys
import base64
import io
from pathlib import Path
from skills.base import BaseSkill
from adapters.llm.vision_client import VisionClient


class VisionSkill(BaseSkill):

    def __init__(self, glm_api_key: str = "", sense_api_key: str = "",
                 llm_client=None, queue=None):
        self._vision = VisionClient(
            glm_key=glm_api_key,
            sense_key=sense_api_key,
            priority=["glm", "sensetime"],
        )
        self._llm = llm_client
        self._queue = queue

    @property
    def name(self) -> str:
        return "vision"

    def get_tools(self) -> list[dict]:
        return [
            {"name": "vision_analyze", "fn": self.vision_analyze,
             "schema": self._s("用 GLM 视觉模型分析图片内容", {
                 "image_b64": {"type": "string", "description": "base64 图片"},
                 "prompt": {"type": "string", "description": "分析指令"},
             }, ["image_b64", "prompt"])},
            {"name": "vision_status", "fn": self.vision_status,
             "schema": self._s("查看视觉任务队列状态", {}, [])},
            {"name": "excel_describe_format", "fn": self.excel_describe_format,
             "schema": self._s("直接读取 Excel 格式信息（字体/颜色/边框/合并），无需截图", {
                 "filepath": {"type": "string", "description": "文件路径"},
                 "sheet": {"type": "string", "description": "Sheet 名称"},
                 "range": {"type": "string", "description": "区域如 A1:G20，默认已用区域"},
             }, ["filepath"])},
            {"name": "excel_understand", "fn": self.excel_understand,
             "schema": self._s("首次打开 Excel 时调用——返回自然语言的使用说明书（表头含义/下拉限制/公式列/空行位置/怎么新增数据）。后续操作直接参考此结果，不用重复调 describe。", {
                 "filepath": {"type": "string", "description": "文件路径"},
             }, ["filepath"])},
            {"name": "word_describe_format", "fn": self.word_describe_format,
             "schema": self._s("直接读取 Word 文档格式信息（字体/段落样式/表格结构），无需截图", {
                 "filepath": {"type": "string", "description": "文件路径"},
             }, ["filepath"])},
            {"name": "pdf_read", "fn": self.pdf_read,
             "schema": self._s("读取 PDF 文件文本内容（支持简历、合同、报表）。超长 PDF 用 start_page 分段读，不要一次读全部", {
                 "filepath": {"type": "string", "description": "PDF 文件路径"},
                 "max_pages": {"type": "integer", "description": "本次读取页数，默认10"},
                 "start_page": {"type": "integer", "description": "起始页码（1开始），默认1"},
             }, ["filepath"])},
            {"name": "pdf_analyze", "fn": self.pdf_analyze,
             "schema": self._s("用 GLM 视觉分析 PDF 单页（扫描件/图片型 PDF）", {
                 "filepath": {"type": "string", "description": "PDF 文件路径"},
                 "prompt": {"type": "string", "description": "分析指令"},
                 "page": {"type": "integer", "description": "页码，默认0=全部"},
             }, ["filepath"])},
            {"name": "pdf_analyze_range", "fn": self.pdf_analyze_range,
             "schema": self._s("批量视觉分析 PDF 页码范围（扫描件专用）——所有页入队后台处理，返回任务ID列表。之后用 vision_get_result 逐个取结果。不要逐页调用 pdf_analyze", {
                 "filepath": {"type": "string", "description": "PDF 文件路径"},
                 "start_page": {"type": "integer", "description": "起始页码（1开始）"},
                 "end_page": {"type": "integer", "description": "结束页码（含）"},
                 "prompt": {"type": "string", "description": "分析指令，如'提取本页所有文字内容'"},
             }, ["filepath", "start_page", "end_page"])},
            {"name": "vision_get_result", "fn": self.vision_get_result,
             "schema": self._s("按任务ID获取视觉分析结果（pdf_analyze_range 返回的 ID）——done 返回内容，pending 返回等待提示", {
                 "task_id": {"type": "string", "description": "任务ID（vq_ 开头）"},
             }, ["task_id"])},
            {"name": "vision_status", "fn": self.vision_status,
             "schema": self._s("查看视觉任务队列状态", {}, [])},
            {"name": "pdf_split", "fn": self.pdf_split,
             "schema": self._s("切割 PDF：提取指定页面范围为新文件", {
                 "filepath": {"type": "string", "description": "PDF 文件路径"},
                 "output": {"type": "string", "description": "输出文件路径"},
                 "start_page": {"type": "integer", "description": "起始页码（1开始）"},
                 "end_page": {"type": "integer", "description": "结束页码"},
             }, ["filepath", "output", "start_page", "end_page"])},
            {"name": "pdf_merge", "fn": self.pdf_merge,
             "schema": self._s("合并多个 PDF 为一个文件", {
                 "files": {"type": "array", "items": {"type": "string"}, "description": "PDF 文件路径列表"},
                 "output": {"type": "string", "description": "输出文件路径"},
             }, ["files", "output"])},
        ]

    # ── 截图 ──

    def screenshot_excel(self, filepath: str, range: str = "") -> str:
        """win32com 截取 Excel 区域。"""
        import win32com.client
        try:
            excel = win32com.client.GetActiveObject("Excel.Application")
        except Exception:
            try:
                excel = win32com.client.Dispatch("Excel.Application")
                excel.Visible = False
            except Exception:
                return _screenshot_fallback(filepath, "excel")

        path = Path(filepath).resolve()
        wb = None
        for w in excel.Workbooks:
            try:
                if Path(w.FullName).resolve() == path:
                    wb = w; break
            except Exception:
                continue
        if not wb:
            # 尝试后台打开
            try:
                wb = excel.Workbooks.Open(str(path))
            except Exception:
                return _screenshot_fallback(filepath, "excel")

        ws = wb.ActiveSheet
        rng = ws.Range(range) if range else ws.UsedRange
        rng.CopyPicture(Format=2)  # xlBitmap

        return _clipboard_to_base64()

    def screenshot_word(self, filepath: str) -> str:
        """截取 Word 页面。"""
        import win32com.client
        try:
            word = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
            except Exception:
                return _screenshot_fallback(filepath, "word")

        path = Path(filepath).resolve()
        doc = None
        for d in word.Documents:
            try:
                if Path(d.FullName).resolve() == path:
                    doc = d; break
            except Exception:
                continue
        if not doc:
            try:
                doc = word.Documents.Open(str(path))
            except Exception:
                return _screenshot_fallback(filepath, "word")

        # 导出为 PDF 单页 → 转图
        import tempfile
        tmp_pdf = Path(tempfile.gettempdir()) / "mother_screenshot.pdf"
        doc.ExportAsFixedFormat(str(tmp_pdf), 17)  # wdFormatPDF

        try:
            import fitz  # PyMuPDF
            pdf = fitz.open(str(tmp_pdf))
            page = pdf[0]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            pdf.close()
            tmp_pdf.unlink(missing_ok=True)
            return base64.b64encode(img_bytes).decode()
        except ImportError:
            tmp_pdf.unlink(missing_ok=True)
            return "⚠️ 需要安装 PyMuPDF: pip install PyMuPDF"

    # ── 分析 ──

    async def vision_analyze(self, image_b64: str, prompt: str) -> str:
        """调视觉模型分析图片。直接同步等待结果（不入队——入队会导致拿不到结果）。"""
        try:
            return await self._vision.analyze(image_b64, prompt)
        except Exception as e:
            return f"❌ 视觉分析失败: {e}"

    def vision_status(self) -> str:
        """查看视觉任务队列状态。"""
        if not self._queue:
            return "（视觉队列未启用）"
        s = self._queue.status()
        lines = [f"📊 视觉队列: {s['done']}完成 {s['pending']}待处理 {s['failed']}失败"]
        if s['pending'] > 0:
            lines.append(f"   后台处理中，预计{s['pending']*3}秒完成")
        if s['failed'] > 0:
            results = self._queue.get_results(5)
            for r in results:
                if r['status'] == 'failed':
                    lines.append(f"   ❌ {r['id'][:12]}: {r['error'][:80]}")
        return "\n".join(lines)

    def excel_describe_format(self, filepath: str, sheet: str = "", range: str = "") -> str:
        """读取 Excel 的完整结构：表头/列名/数据验证/公式/合并/冻结/空行/行数。"""
        path = Path(filepath.strip('"').strip("'").strip() if filepath else "")
        if not path.is_absolute():
            for d in ["./workbooks", "./output", "."]:
                p = Path(d) / path.name
                if p.exists(): path = p; break
        if not path.exists():
            return f"❌ 文件不存在: {path.name}"
        try:
            import openpyxl
            from openpyxl.utils import get_column_letter
            wb = openpyxl.load_workbook(str(path))
            ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
            lines = []

            lines.append(f"📋 文件: {path.name}")
            lines.append(f"Sheet: {ws.title}（共 {ws.max_row} 行，{get_column_letter(ws.max_column)} 共{ws.max_column}列）")
            lines.append("")

            # ── 表头（第1行）──
            header_cells = list(ws.iter_rows(min_row=1, max_row=1, max_col=min(ws.max_column, 20), values_only=True))[0]
            header_items = [f"{get_column_letter(i+1)}={v}" for i, v in enumerate(header_cells) if v is not None]
            lines.append(f"表头（第1行）: {'  '.join(header_items)}" if header_items else "表头: （空）")
            lines.append("")

            # ── 数据验证（下拉）──
            if ws.data_validations and ws.data_validations.dataValidation:
                for dv in list(ws.data_validations.dataValidation)[:10]:
                    rng = str(dv.sqref) if hasattr(dv, 'sqref') else ""
                    formula = str(dv.formula1) if hasattr(dv, 'formula1') else ""
                    if formula.startswith('"'):
                        items = formula.strip('"').split(",")
                        formula = ", ".join(i.strip() for i in items)
                    lines.append(f"  📌 {rng}: 下拉选项 → {formula}" if rng else f"  📌 下拉: {formula}")
                if lines[-1] != "":
                    lines.append("")

            # ── 公式列 ──
            formula_cols = {}
            for row in ws.iter_rows(min_row=2, max_row=min(ws.max_row, 5), max_col=ws.max_column):
                for cell in row:
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        col = get_column_letter(cell.column)
                        if col not in formula_cols:
                            formula_cols[col] = str(cell.value)[:60]
            if formula_cols:
                for col, fml in sorted(formula_cols.items()):
                    lines.append(f"  🔢 {col}列: {fml}")
                lines.append("")

            # ── 合并单元格 ──
            if ws.merged_cells.ranges:
                merges = [str(m) for m in list(ws.merged_cells.ranges)[:10]]
                lines.append(f"合并单元格: {', '.join(merges)}")
                lines.append("")

            # ── 冻结窗格 ──
            if ws.freeze_panes:
                lines.append(f"冻结窗格: {ws.freeze_panes}")
                lines.append("")

            # ── 空行分隔 ──
            empty_rows = []
            for row in ws.iter_rows(min_row=2, max_row=min(ws.max_row, 100)):
                vals = [c.value for c in row]
                if all(v is None for v in vals):
                    empty_rows.append(row[0].row)
            if empty_rows:
                lines.append(f"空行: 第{', '.join(str(r) for r in empty_rows[:8])}行" + 
                           (f" 等{len(empty_rows)}处" if len(empty_rows) > 8 else ""))
                lines.append("")

            # ── 前5行预览 ──
            lines.append(f"前{min(5, ws.max_row)}行预览:")
            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 6), max_col=min(ws.max_column, 8), values_only=True):
                lines.append("  " + " | ".join(str(v)[:20] if v is not None else "" for v in row))

            wb.close()
            return "\n".join(lines)[:3000]
        except Exception as e:
            return f"❌ 读取格式失败: {e}"

    async def excel_understand(self, filepath: str) -> str:
        """首次打开 Excel → LLM 生成自然语言使用说明书。"""
        # 先拿结构数据
        raw = self.excel_describe_format(filepath)
        if raw.startswith("❌"):
            return raw
        # 调 LLM 理解
        if not self._llm:
            return raw + "\n\n（提示：未配置 LLM，上述为原始结构数据）"
        try:
            prompt_path = Path(__file__).parent.parent / "prompts" / "understand_excel.txt"
            sys_prompt = open(str(prompt_path), "r", encoding="utf-8").read()
            messages = [
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": f"以下是一个 Excel 文件的结构描述，请按规范输出使用说明书：\n\n{raw}"},
            ]
            resp = await self._llm.chat(messages, temperature=0.2)
            return f"📋 文件理解（LLM）:\n{resp.content}\n\n（后续操作请据此说明书进行，无需重复调 understand）"
        except Exception as e:
            return raw + f"\n\n（理解失败: {e}，请手动分析上述结构）"

    def word_describe_format(self, filepath: str) -> str:
        """python-docx 直接读 Word 格式。"""
        from skills.word_skill import WordSkill
        try:
            path = WordSkill._ensure_docx(filepath)
        except Exception as e:
            return str(e)
        try:
            from docx import Document
            doc = Document(str(path))
            lines = []
            for i, para in enumerate(doc.paragraphs[:15]):
                if not para.text.strip():
                    continue
                info = f"  [{para.style.name if para.style else ''}] {para.text[:60]}"
                if para.runs:
                    r = para.runs[0]
                    a = []
                    if r.font.bold: a.append("加粗")
                    if r.font.size: a.append(f"{r.font.size.pt}pt")
                    if a:
                        info += f" ({', '.join(a)})"
                lines.append(info)
            for ti, table in enumerate(doc.tables):
                lines.append(f"\n表格{ti + 1}: {len(table.rows)}行×{len(table.columns)}列")
                for ri, row in enumerate(table.rows[:4]):
                    cells = [c.text.strip()[:18] for c in row.cells]
                    lines.append(f"  {', '.join(cells)}")
            return "\n".join(lines)[:2500]
        except Exception as e:
            return f"❌ 读取失败: {e}"

    def pdf_read(self, filepath: str, max_pages: int = 10, start_page: int = 1) -> str:
        """PyMuPDF 提取 PDF 文本。start_page 起读，支持长 PDF 分段读取。"""
        path = Path(filepath)
        if not path.is_absolute():
            for d in ["./workbooks", "./output", "."]:
                p = Path(d) / path.name
                if p.exists(): path = p; break
        if not path.exists():
            return f"❌ 文件不存在: {filepath}"
        try:
            import fitz
            doc = fitz.open(str(path))
            total = doc.page_count
            start = max(1, int(start_page or 1))
            end = min(total, start + max_pages - 1)
            lines = [f"PDF: {path.name} | {total}页（本次读第{start}-{end}页）"]
            for i in range(start - 1, end):
                page = doc[i]
                text = page.get_text().strip()
                if text:
                    lines.append(f"\n--- 第{i+1}页 ---")
                    lines.append(text[:3000])
                else:
                    lines.append(f"\n--- 第{i+1}页（无文本层，需用 pdf_analyze_range 视觉分析）---")
            doc.close()
            result = "\n".join(lines)
            # 超长截断，但保留页标记
            if len(result) > 15000:
                result = result[:15000] + f"\n…（已截断，可用 start_page 续读第{start + 5}页起）"
            return result
        except Exception as e:
            return f"❌ PDF读取失败: {e}"

    async def pdf_analyze(self, filepath: str, prompt: str = "", page: int = 0) -> str:
        """PDF 页面 → 图片 → 视觉分析。"""
        path = Path(filepath)
        if not path.is_absolute():
            for d in ["./workbooks", "./output", "."]:
                p = Path(d) / path.name
                if p.exists(): path = p; break
        if not path.exists():
            return f"❌ 文件不存在: {filepath}"
        try:
            import fitz
            doc = fitz.open(str(path))
            p = page - 1 if page > 0 else 0
            if p >= doc.page_count:
                doc.close(); return f"❌ 页码超出范围（共{doc.page_count}页）"
            pix = doc[p].get_pixmap(dpi=150)
            img_b64 = base64.b64encode(pix.tobytes("png")).decode()
            doc.close()
            return await self.vision_analyze(img_b64, prompt or f"分析这个PDF页面的内容和布局")
        except Exception as e:
            return f"❌ PDF分析失败: {e}"

    async def pdf_analyze_range(self, filepath: str, start_page: int, end_page: int,
                                 prompt: str = "") -> str:
        """批量视觉分析 PDF 页码范围——全部入队，返回任务 ID 列表。"""
        path = Path(filepath)
        if not path.is_absolute():
            for d in ["./workbooks", "./output", "."]:
                p = Path(d) / path.name
                if p.exists(): path = p; break
        if not path.exists():
            return f"❌ 文件不存在: {filepath}"
        if not self._queue:
            return "❌ 视觉队列未启用"
        try:
            import fitz
            doc = fitz.open(str(path))
            total = doc.page_count
            s = max(1, int(start_page)); e = min(total, int(end_page))
            if s > e or s > total:
                doc.close(); return f"❌ 页码范围错误（共{total}页）"
            tids = []
            for i in range(s - 1, e):
                pix = doc[i].get_pixmap(dpi=150)
                img_b64 = base64.b64encode(pix.tobytes("png")).decode()
                tid = self._queue.add(img_b64, prompt or f"提取PDF第{i+1}页的全部文字内容")
                tids.append((i + 1, tid))
            doc.close()
            lines = [f"📬 已入队 {len(tids)} 页（第{s}-{e}页），任务ID列表:"]
            for pg, tid in tids:
                lines.append(f"  第{pg}页: {tid[:12]}")
            lines.append(f"预计 {len(tids)*3} 秒完成。用 vision_get_result(任务ID) 取结果，一次一页。")
            return "\n".join(lines)
        except Exception as ex:
            return f"❌ 批量分析失败: {ex}"

    def vision_get_result(self, task_id: str) -> str:
        """按任务 ID 取视觉分析结果。"""
        if not self._queue:
            return "❌ 视觉队列未启用"
        r = self._queue.get_result(task_id.strip())
        if not r:
            return f"❌ 任务不存在: {task_id}（请确认 ID 完整，或以 vq_ 开头）"
        if r["status"] == "done":
            return f"✅ 第{task_id[:12]}页分析完成:\n{r['result']}"
        if r["status"] == "failed":
            return f"❌ 第{task_id[:12]}页分析失败: {r['error']}"
        return f"⏳ 第{task_id[:12]}页还在处理中（{r['status']}），稍后再查"

    def pdf_split(self, filepath: str, output: str, start_page: int, end_page: int) -> str:
        path = Path(filepath)
        if not path.exists():
            return f"❌ 文件不存在: {filepath}"
        try:
            import fitz
            doc = fitz.open(str(path))
            if start_page < 1 or end_page > doc.page_count:
                doc.close(); return f"❌ 页码范围错误（共{doc.page_count}页）"
            new = fitz.open()
            new.insert_pdf(doc, from_page=start_page-1, to_page=end_page-1)
            Path(output).parent.mkdir(parents=True, exist_ok=True)
            new.save(output); new.close(); doc.close()
            return f"✅ 已切割: 第{start_page}-{end_page}页 → {Path(output).name}"
        except Exception as e:
            return f"❌ 切割失败: {e}"

    def pdf_merge(self, files: list, output: str) -> str:
        try:
            import fitz
            merged = fitz.open()
            for f in files:
                p = Path(f)
                if not p.exists(): return f"❌ 文件不存在: {f}"
                doc = fitz.open(str(p)); merged.insert_pdf(doc); doc.close()
            Path(output).parent.mkdir(parents=True, exist_ok=True)
            merged.save(output); merged.close()
            return f"✅ 已合并 {len(files)} 个 PDF → {Path(output).name}"
        except Exception as e:
            return f"❌ 合并失败: {e}"


def _screenshot_fallback(filepath: str, app: str) -> str:
    """截图降级：pywin32 不可用。"""
    return f"⚠️ 无法截取 {app} 截图（pywin32 未安装或文件未打开）。请先打开文件再试。"


def _clipboard_to_base64() -> str:
    """从剪贴板读取图片 → base64。"""
    try:
        from PIL import ImageGrab
        img = ImageGrab.grabclipboard()
        if img:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return base64.b64encode(buf.getvalue()).decode()
        return "⚠️ 剪贴板中没有图片"
    except Exception as e:
        return f"⚠️ 截图失败: {e}"
