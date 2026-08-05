"""
Word 专业文档渲染器单元测试。

运行: python -m unittest tests.test_word_renderer -v
覆盖:
    - 标题层级（Heading 0/1/2）
    - 目录字段（TOC fldChar）
    - 中文字体（eastAsia）
    - 首行缩进
    - 文档信息表格 / 数据表格
    - 无效 JSON 错误处理
"""
import unittest
import sys
import json
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from adapters.word_doc_renderer import render_document_json, render_document_json_str


class TestDocumentRenderer(unittest.TestCase):
    def setUp(self):
        self.tmpdir = Path(tempfile.mkdtemp())
        self.out = self.tmpdir / "test.docx"

    def _sample(self):
        return {
            "title": "测试文档",
            "meta": [["编号", "T-001"]],
            "toc": True,
            "sections": [
                {"level": 1, "title": "1.0 章节",
                 "paragraphs": ["正文一。", "正文二。"],
                 "bullets": ["项目A", "项目B"]},
                {"level": 1, "title": "2.0 子节",
                 "sections": [{"level": 2, "title": "2.1 小节",
                               "paragraphs": ["子节内容。"],
                               "numbered": ["步骤1", "步骤2"]}]},
            ],
        }

    def test_render_success(self):
        r = render_document_json(self._sample(), str(self.out))
        self.assertIn("✅", r)
        self.assertTrue(self.out.exists())

    def test_heading_structure(self):
        from docx import Document
        render_document_json(self._sample(), str(self.out))
        doc = Document(str(self.out))
        headings = [(p.style.name, p.text) for p in doc.paragraphs
                    if p.style.name.startswith("Heading") or p.style.name == "Title"]
        self.assertEqual(headings[0], ("Title", "测试文档"))
        self.assertIn(("Heading 1", "1.0 章节"), headings)
        self.assertIn(("Heading 1", "2.0 子节"), headings)
        self.assertIn(("Heading 2", "2.1 小节"), headings)

    def test_toc_field(self):
        from zipfile import ZipFile
        render_document_json(self._sample(), str(self.out))
        with ZipFile(self.out) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        self.assertIn("TOC", xml)
        self.assertIn("fldChar", xml)

    def test_east_asia_font(self):
        from zipfile import ZipFile
        render_document_json(self._sample(), str(self.out))
        with ZipFile(self.out) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        self.assertIn("w:eastAsia", xml)

    def test_first_line_indent(self):
        from zipfile import ZipFile
        render_document_json(self._sample(), str(self.out))
        with ZipFile(self.out) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        self.assertIn("w:firstLine", xml)

    def test_tables(self):
        from docx import Document
        render_document_json(self._sample(), str(self.out))
        doc = Document(str(self.out))
        self.assertGreaterEqual(len(doc.tables), 1)  # meta 表

    def test_invalid_json(self):
        r = render_document_json_str("{broken json", str(self.out))
        self.assertIn("❌", r)

    def test_empty_sections(self):
        r = render_document_json({"title": "空文档"}, str(self.out))
        self.assertIn("✅", r)


if __name__ == "__main__":
    unittest.main()
